import re

from babel.dates import format_date
from babel.numbers import format_currency
from django.utils import timezone
from docx import Document
from docxtpl import DocxTemplate
from num2words import num2words

from gelfmp.models.choices import ContactType
from gelfmp.templatetags.custom_filters import format_cpf_cnpj
from gelfmp.utils.normalization import normalize_cep, normalize_name


class ContractFillerService:
    TEMPLATE_PATH = 'static/data/base_charcoal_contract.docx'
    PRICE_DENSITY = 0.235

    def fill_contract(self, supplier, contract, output, context=None):
        doc = DocxTemplate(self.TEMPLATE_PATH)

        context = context or self.build_context(supplier, contract)
        self.validate_context(context)

        doc.render(context)
        doc.save(output)

    def build_context(self, supplier, contract):
        contacts_context = self.get_contacts_context(supplier)
        price_volume_context = self.get_price_volume_context(contract)
        bank_context = self.get_bank_context(supplier)

        corporate_name = self.get_corporate_name(supplier.corporate_name)
        footer_corporate_name = normalize_name(corporate_name)
        cpf_cnpj = format_cpf_cnpj(supplier.cpf_cnpj)

        city_and_state = f'{supplier.city.name}/{supplier.state.abbr}'
        dcf = contract.dcf.process_number

        if contract.entry_date:
            entry_date_formatted = contract.entry_date.strftime('%d/%m/%Y')
        else:
            entry_date_formatted = None

        return {
            'header_id': self.get_header_id(supplier, contract),
            'footer_corporate_name': footer_corporate_name,
            'corporate_name': corporate_name,
            'farm': self.get_farm_name(supplier),
            'city_and_state': city_and_state,
            'cep': normalize_cep(supplier.cep),
            'cpf_cnpj': cpf_cnpj,
            'dcf': dcf,
            'entry_date': entry_date_formatted,
            'today': self.get_formatted_date(timezone.now()),
            **bank_context,
            **price_volume_context,
            **contacts_context,
        }

    def get_header_id(self, supplier, contract):
        return f'{contract.id:03}_{contract.entry_date.year}_CV_{self.normalize_header(supplier.corporate_name)}'.ljust(
            200  # Tamanho da linha do cabeçalho para alinhar a logo da GELF a direita da página.
        )

    def get_farm_name(self, supplier):
        found = re.search(r'\s+FAZENDA\s+(.+)', supplier.corporate_name, re.IGNORECASE)
        if found:
            return normalize_name(found.group(1))

        raise ValueError(f'Fazenda não encontrada no nome do fornecedor: {supplier.corporate_name}.')

    def get_corporate_name(self, corporate_name):
        splited = corporate_name.split('-')

        if len(splited) > 1:
            return ' '.join(splited[:-1]).strip()
        else:
            return corporate_name.strip()

    def get_formatted_date(self, date):
        return format_date(date, format="d 'de' MMMM 'de' yyyy", locale='pt_BR')

    def get_price_volume_context(self, contract):
        contract_volume = contract.contract_volume
        contract_volume_in_words = num2words(contract_volume, lang='pt_BR')

        price = contract.price
        price_per_ton = round(price / self.PRICE_DENSITY, 2)
        price_formatted = format_currency(price, 'BRL', locale='pt_BR')
        price_per_ton_formatted = format_currency(price_per_ton, 'BRL', locale='pt_BR')

        estimated_total = round(contract_volume * price, 2)
        estimated_total_formatted = format_currency(estimated_total, 'BRL', locale='pt_BR')
        estimated_total_in_words = num2words(estimated_total, lang='pt_BR') + ' reais'

        return {
            'contract_volume': contract_volume,
            'contract_volume_in_words': contract_volume_in_words,
            'price': price_formatted,
            'price_per_ton': price_per_ton_formatted,
            'estimated_total': estimated_total_formatted,
            'estimated_total_in_words': estimated_total_in_words,
        }

    def get_contacts_context(self, supplier):
        witness = supplier.contacts.filter(contact_type=ContactType.WITNESS).first()
        legal_representatives = supplier.contacts.filter(contact_type=ContactType.LEGAL_REPRESENTATIVE).all()[:2]

        if not witness or not legal_representatives.exists():
            raise ValueError('Testemunha ou representante legal não encontrado.')

        return {
            'witness': witness.name,
            'witness_email': witness.email,
            'witness_cpf': format_cpf_cnpj(witness.cpf),
            'legal_representative': legal_representatives[0].name,
            'legal_representative_email': legal_representatives[0].email,
            'legal_representative_cpf': format_cpf_cnpj(legal_representatives[0].cpf),
            'legal_representative2': None if len(legal_representatives) < 2 else legal_representatives[1].name,
            'legal_representative2_email': None if len(legal_representatives) < 2 else legal_representatives[1].email,
            'legal_representative2_cpf': None
            if len(legal_representatives) < 2
            else format_cpf_cnpj(legal_representatives[1].cpf),
        }

    def get_bank_context(self, supplier):
        return {
            'bank_name': supplier.bank_details.bank_name,
            'bank_account_number': supplier.bank_details.account_number,
            'bank_agency': supplier.bank_details.agency,
        }

    def normalize_header(self, text):
        text = re.sub(r'[^a-zA-Z0-9]', '_', text)
        text = re.sub(r'_{2,}', '_', text)

        return text

    def validate_context(self, context):
        variables = set()

        doc = Document(self.TEMPLATE_PATH)
        placeholder_regex = r'\{\{\s*(.+?)\s*\}\}'

        for paragraph in doc.paragraphs:
            variables.update(re.findall(placeholder_regex, paragraph.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        variables.update(re.findall(placeholder_regex, paragraph.text))

        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    variables.update(re.findall(placeholder_regex, paragraph.text))

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    variables.update(re.findall(placeholder_regex, paragraph.text))

        missing_variables = variables - set(context.keys())

        if missing_variables:
            raise ValueError(f'Os seguintes campos do contexto faltam: {", ".join(missing_variables)}')
